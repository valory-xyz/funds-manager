"""Generate a management-friendly Working Document (.docx) from the tech spec."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_table_row(table, cells, bold=False, header=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "2F5496")
            run.font.color.rgb = RGBColor(255, 255, 255)
    return row


def style_table(table):
    """Apply consistent styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)


doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Default font --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Auto-Transfer Safe → EOA\nfor Native Gas Funding')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(47, 84, 150)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Working Document')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Status: Draft\nDate: 2026-03-18\nRepo: valory-xyz/funds-manager')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Problem Statement',
    '2. Current Architecture',
    '3. Proposed Solution',
    '4. How It Works — End-to-End',
    '5. Impact on Each Agent',
    '6. What Changes, What Doesn\'t',
    '7. Configuration Changes',
    '8. Implementation Plan',
    '9. Edge Cases & Mitigations',
    '10. Decisions Made',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. PROBLEM STATEMENT
# ============================================================
doc.add_heading('1. Problem Statement', level=1)

doc.add_paragraph(
    "Pearl's agent wallet uses two separate addresses under the hood:"
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
add_table_row(table, ['', 'Agent EOA', 'Agent Safe'], header=True)
# remove auto-created first row
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['', 'Agent EOA', 'Agent Safe'], header=True)
add_table_row(table, ['Purpose', 'Gas wallet — submits transactions', 'Funds wallet — holds operating capital'])
add_table_row(table, ['Funded by', 'Middleware (currently)', 'User via Pearl UI'])
add_table_row(table, ['Visible to user?', 'No (internal)', 'Yes ("the wallet")'])
style_table(table)

doc.add_paragraph('')
doc.add_heading('What the User Experiences', level=2)

items = [
    'The Agent EOA drains to ~0 native tokens (e.g., POL on Polygon). This is accelerated by x402 payment usage.',
    'Meanwhile, the Agent Safe still holds native tokens — e.g., 50 POL.',
    'Pearl shows an alert: "Fund your agent with 30 POL".',
    'The user opens their wallet and sees 50 POL already there.',
    'Result: The user believes they\'ve already funded the agent. The agent appears "running" but cannot execute any transactions.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Root Cause', level=2)
doc.add_paragraph(
    'The EOA and Safe operate independently. When the EOA runs out of gas, '
    'it has no mechanism to draw from the Safe\'s funds. The user is asked '
    'to provide new funds even though sufficient funds already exist in the Safe.'
)

doc.add_heading('Goal', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'When the Agent EOA is low on native gas, the agent should automatically '
    'transfer native tokens from Agent Safe → Agent EOA, without user intervention.'
)
run.bold = True

doc.add_page_break()

# ============================================================
# 2. CURRENT ARCHITECTURE
# ============================================================
doc.add_heading('2. Current Architecture', level=1)

doc.add_heading('Current Funding Flow', level=2)

steps = [
    ('Agent', 'The funds_manager skill queries on-chain balances via multicall. Computes deficit (topup − balance) for each token. Exposes results via the /funds-status HTTP endpoint.'),
    ('Middleware', 'Periodically polls /funds-status. If deficits exist, passes them to the frontend via /api/v2/service/{id}/funding_requirements.'),
    ('User', 'Sees a funding alert on the Pearl UI. Optionally adjusts amounts and approves.'),
    ('Frontend', 'Calls the middleware /fund endpoint with user-approved amounts.'),
    ('Middleware', 'Transfers from MasterSafe → Agent EOA and/or Agent Safe.'),
]
for actor, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{actor}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Agents Using This Skill', level=2)

doc.add_paragraph(
    'The funds_manager skill is consumed by four agents across four chains:'
)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
add_table_row(table, ['Agent', 'Chain', 'EOA Topup', 'Safe Topup', 'Safe Enabled?'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Agent', 'Chain', 'EOA Topup', 'Safe Topup', 'Safe Enabled?'], header=True)
add_table_row(table, ['Optimus', 'Optimism', '0.0002 ETH', '0 (disabled)', 'No'])
add_table_row(table, ['Omenstrat', 'Gnosis', '2.0 xDAI', '5.0 xDAI', 'Yes'])
add_table_row(table, ['Polystrat', 'Polygon', '30.0 POL', '40.0 POL', 'Yes'])
add_table_row(table, ['Agents Fun', 'Base', '~0.000326 ETH', '~0.001629 ETH', 'Yes'])
style_table(table)

doc.add_paragraph('')

doc.add_heading('Per-Agent Customizations (Unchanged)', level=2)
doc.add_paragraph(
    'Each agent\'s HTTP handler applies additional logic after calling the '
    'funds_manager skill. These customizations remain in each agent\'s '
    'codebase and are not affected by this proposal. There are two types:'
)

p = doc.add_paragraph()
run = p.add_run('Balance Adjustments (read-only, no on-chain action):')
run.bold = True
items = [
    'Omenstrat: wxDAI balance folded into native xDAI (1:1 peg) for Safe deficit calculation.',
    'Polystrat: USDC balance converted to POL-equivalent via CoinGecko exchange rate, folded into POL for Safe deficit calculation.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('x402 USDC Swap (on-chain EOA transaction, all agents):')
run.bold = True
items = [
    'All four agents support x402 payments. When enabled, each /funds-status poll triggers an ETH→USDC swap via LiFi in a background thread, executed from the EOA (not the Safe).',
    'If the swap fails, Optimus, Omenstrat, and Polystrat inject the ETH cost into the /funds-status response so MW funds the EOA with extra ETH. Agents Fun does not inject on failure.',
    'Optimus also has withdrawal mode logic that overrides the /funds-status response with minimal per-action gas deficits.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. PROPOSED SOLUTION
# ============================================================
doc.add_heading('3. Proposed Solution', level=1)

doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'Add an auto-transfer capability to the funds_manager skill. When the Agent EOA\'s '
    'native token balance drops below a configured threshold, a new FSM round automatically '
    'transfers native tokens from the Agent Safe to the Agent EOA — without any user action.'
)

doc.add_heading('Two Independent Systems, One Skill', level=2)
doc.add_paragraph(
    'The solution adds a new FSM component alongside the existing balance-checking behaviour. '
    'They operate independently within the same skill:'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
add_table_row(table, ['Component', 'Purpose', 'Changes?'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Component', 'Purpose', 'Changes?'], header=True)
add_table_row(table, [
    'FundsManagerBehaviour\n(existing)',
    'Reads balances, computes deficits, serves /funds-status to middleware',
    'No code changes.\nOnly config changes (EOA threshold lowered).'
])
add_table_row(table, [
    'FundsManagerAbciApp\n(new)',
    'Checks EOA balance each FSM cycle, auto-transfers from Safe → EOA when needed',
    'New. Lives in funds-manager repo.\nAgents import and compose into their FSM.'
])
style_table(table)

doc.add_paragraph('')

doc.add_heading('Three Funding Paths (Prioritized)', level=2)

p = doc.add_paragraph()
run = p.add_run('Path 1 — Auto-Transfer (normal case, no user action)')
run.bold = True
doc.add_paragraph(
    'At the start of each FSM cycle, the new round checks the EOA balance. '
    'If it\'s below the auto-transfer threshold and the Safe has native funds, '
    'the agent executes a Safe→EOA transfer. The user sees nothing — the agent '
    'self-funds from its own Safe.',
    style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run('Path 2 — User Funds Safe (Safe is depleted)')
run.bold = True
doc.add_paragraph(
    'If auto-transfers deplete the Safe below its own threshold, the existing '
    '/funds-status endpoint reports a Safe deficit. The middleware alerts the user: '
    '"Fund your Safe with X." The user funds the Safe, and auto-transfer resumes '
    'on the next cycle.',
    style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run('Path 3 — MW Funds EOA Directly (last resort)')
run.bold = True
doc.add_paragraph(
    'If the EOA is so low it cannot even execute a Safe transaction (below safe_tx_value), '
    'the existing /funds-status endpoint reports an EOA deficit. The middleware funds the '
    'EOA directly. This should be rare with correct threshold configuration.',
    style='List Bullet'
)

doc.add_page_break()

# ============================================================
# 4. HOW IT WORKS — END TO END
# ============================================================
doc.add_heading('4. How It Works — End-to-End', level=1)

doc.add_paragraph(
    'This section traces the complete lifecycle using Omenstrat (Gnosis) as a concrete example. '
    'EOA has 0.1 xDAI (below 0.21 threshold). Safe has 3.0 xDAI.'
)

doc.add_heading('Step 1: FSM Round Auto-Transfers', level=2)
steps = [
    'Agent FSM cycle starts. FundsManagerAbciApp runs first.',
    'CheckFundsRound reads balances: EOA = 0.1 xDAI, Safe = 3.0 xDAI.',
    'EOA (0.1) < auto_transfer_threshold (0.21) → transfer needed.',
    'Transfer amount = min(3.0, 2.0 − 0.1) = 1.9 xDAI.',
    'PrepareSafeToEOATransferRound builds the Safe transaction.',
    'TransactionSettlementAbci signs, submits, and confirms on-chain.',
    'Result: EOA ≈ 2.0 xDAI, Safe ≈ 1.1 xDAI. No user action required.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{step}', style='List Number')

doc.add_heading('Step 2: Middleware Polls (Asynchronous)', level=2)
steps = [
    'Middleware polls GET /funds-status (unchanged code).',
    'funds_manager reads on-chain balances: EOA = 2.0, Safe = 1.1.',
    'EOA (2.0) > 0.001 threshold → deficit = 0. Safe (1.1) > 1.0 threshold → deficit = 0.',
    'Handler applies wxDAI adjustment (unchanged). Still no deficit.',
    'Middleware receives: all deficits = 0. Nothing to report.',
    'User sees nothing. Agent operating normally.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{step}', style='List Number')

doc.add_heading('Step 3: When Safe Gets Depleted', level=2)
doc.add_paragraph(
    'After several auto-transfers, the Safe\'s native balance drops below its threshold:'
)
steps = [
    'EOA = 0.1 xDAI, Safe = 0.3 xDAI. FSM transfers all 0.3 to EOA.',
    'EOA ≈ 0.4, Safe ≈ 0. Safe (0) < 1.0 threshold → deficit = 5.0 xDAI.',
    'Middleware reports Safe deficit. Pearl shows: "Fund your Safe with 5.0 xDAI."',
    'User approves. Middleware transfers 5.0 xDAI from MasterSafe → Agent Safe.',
    'Next FSM cycle: if EOA is low, auto-transfer from newly funded Safe resumes.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{step}', style='List Number')

doc.add_heading('Key Insight', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'The user always funds the Safe — the wallet they understand. '
    'The agent manages the EOA internally. The user never sees confusing '
    'EOA funding prompts when the Safe has funds.'
)
run.italic = True

doc.add_page_break()

# ============================================================
# 5. IMPACT ON EACH AGENT
# ============================================================
doc.add_heading('5. Impact on Each Agent', level=1)

doc.add_paragraph(
    'All four agents support x402 payments (ETH→USDC swap via LiFi on the EOA, '
    'triggered in a background thread on each /funds-status poll). This is an existing '
    'per-agent mechanism that stays in each agent\'s handler — it is not affected by '
    'the auto-transfer feature. The auto-transfer and x402 swap both submit EOA transactions, '
    'but they run at different times (FSM cycle start vs MW poll) so nonce conflicts are unlikely.'
)
doc.add_paragraph('')

doc.add_heading('Optimus (Optimism)', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Item', 'Detail'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Item', 'Detail'], header=True)
add_table_row(table, ['FSM change', 'Add FundsManagerAbciApp to FSM composition.'])
add_table_row(table, ['Config change', 'Enable Safe native funding (currently disabled — topup=0). Add auto_transfer config. Lower EOA threshold to safe_tx_value.'])
add_table_row(table, ['x402 interaction', 'x402 ETH→USDC swap runs in background thread on each MW poll. Consumes EOA ETH (gas + swap value), accelerating EOA drain. On swap failure, injects ETH deficit into /funds-status response so MW funds EOA with extra ETH.'])
add_table_row(table, ['Withdrawal mode', 'Unaffected — handler logic runs after get_funds_status(), which is unchanged.'])
style_table(table)

doc.add_paragraph('')
doc.add_heading('Omenstrat (Gnosis)', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Item', 'Detail'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Item', 'Detail'], header=True)
add_table_row(table, ['FSM change', 'Add FundsManagerAbciApp to FSM composition.'])
add_table_row(table, ['Config change', 'Add auto_transfer config. Lower EOA threshold to safe_tx_value.'])
add_table_row(table, ['wxDAI adjustment', 'No conflict. Auto-transfer only moves native xDAI, not wxDAI. Handler\'s wxDAI folding is unchanged.'])
add_table_row(table, ['x402 interaction', 'Same as Optimus — shared handler code. ETH→USDC swap in background thread. Same deficit injection on failure.'])
style_table(table)

doc.add_paragraph('')
doc.add_heading('Polystrat (Polygon)', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Item', 'Detail'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Item', 'Detail'], header=True)
add_table_row(table, ['FSM change', 'Add FundsManagerAbciApp to FSM composition.'])
add_table_row(table, ['Config change', 'Add auto_transfer config. Lower EOA threshold to safe_tx_value.'])
add_table_row(table, ['USDC→POL adjustment', 'Known interaction: handler\'s USDC→POL adjustment may mask Safe native deficit after auto-transfer depletes native POL. Requires small handler fix to not suppress native POL deficit when needed for auto-transfer.'])
add_table_row(table, ['x402 interaction', 'Same as Omenstrat — shared handler code. ETH→USDC swap in background thread.'])
style_table(table)

doc.add_paragraph('')
doc.add_heading('Agents Fun / meme-ooorr (Base)', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Item', 'Detail'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Item', 'Detail'], header=True)
add_table_row(table, ['FSM change', 'Add FundsManagerAbciApp to FSM composition, before CheckFundsBehaviour.'])
add_table_row(table, ['Config change', 'Add auto_transfer config. Lower EOA threshold. Consider increasing threshold (currently very close to Safe tx gas cost).'])
add_table_row(table, ['x402 interaction', 'Same background thread swap mechanism. However, meme-ooorr does NOT inject ETH deficit into /funds-status on swap failure (unlike other agents). Swap failure is silent.'])
style_table(table)

doc.add_page_break()

# ============================================================
# 6. WHAT CHANGES, WHAT DOESN'T
# ============================================================
doc.add_heading("6. What Changes, What Doesn't", level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
add_table_row(table, ['Component', 'Changes?', 'Detail'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Component', 'Changes?', 'Detail'], header=True)
add_table_row(table, ['get_funds_status() code', 'No', 'Zero changes to balance checking, deficit calculation, or response format.'])
add_table_row(table, ['/funds-status response', 'No', 'Identical JSON structure. Middleware reads it the same way.'])
add_table_row(table, ['Per-agent HTTP handlers', 'No', 'wxDAI folding, USDC→POL, x402 — all unchanged.'])
add_table_row(table, ['olas-operate-middleware', 'No', 'Reads deficits, forwards to frontend. Zero code changes.'])
add_table_row(table, ['olas-operate-app (frontend)', 'No', 'Displays whatever middleware reports. Zero UI changes.'])
add_table_row(table, ['fund_requirements config', 'Yes', 'EOA native threshold lowered to safe_tx_value per chain.'])
add_table_row(table, ['funds_manager skill', 'Yes', 'New auto_transfer parameters. New FSM rounds, behaviours, payloads.'])
add_table_row(table, ['Per-agent composition.py', 'Yes', 'Each agent imports and composes FundsManagerAbciApp.'])
style_table(table)

doc.add_paragraph('')

doc.add_heading('Scenario Comparison: Before vs After', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
add_table_row(table, ['Scenario', 'Today', 'After', 'Changed?'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Scenario', 'Today', 'After', 'Changed?'], header=True)
add_table_row(table, ['EOA low, Safe has funds', 'User funds EOA', 'Agent auto-transfers (no user action)', 'Yes'])
add_table_row(table, ['EOA low, Safe depleted', 'User funds EOA', 'User funds Safe', 'Yes'])
add_table_row(table, ['EOA low, Safe empty', 'User funds EOA', 'User funds Safe → agent auto-transfers', 'Yes'])
add_table_row(table, ['EOA critically low', 'User funds EOA', 'User funds EOA (fallback)', 'No'])
add_table_row(table, ['Safe ERC20 low', 'User funds Safe', 'User funds Safe', 'No'])
add_table_row(table, ['Safe native low, EOA fine', 'User funds Safe', 'User funds Safe', 'No'])
add_table_row(table, ['Both fine', 'No action', 'No action', 'No'])
style_table(table)

doc.add_page_break()

# ============================================================
# 7. CONFIGURATION CHANGES
# ============================================================
doc.add_heading('7. Configuration Changes', level=1)

doc.add_heading('New Parameter: auto_transfer', level=2)
doc.add_paragraph(
    'A new auto_transfer configuration is added to the funds_manager skill\'s Params model. '
    'This drives the FSM round and is configured per chain:'
)

items = [
    'threshold — EOA balance below which auto-transfer triggers (e.g., 0.21 xDAI for Omenstrat).',
    'topup — target EOA balance after auto-transfer (e.g., 2.0 xDAI for Omenstrat).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'These values are what is currently configured as the EOA threshold/topup in '
    'fund_requirements. They move to the new auto_transfer config.'
)

doc.add_heading('Reconfigured: EOA Threshold in fund_requirements', level=2)
doc.add_paragraph(
    'The existing EOA native threshold in fund_requirements is lowered to safe_tx_value — '
    'the minimum balance needed to execute a Safe transaction. This ensures /funds-status '
    'only reports an EOA deficit as a last resort (Path 3).'
)

doc.add_heading('Example: Omenstrat (Gnosis)', level=2)

p = doc.add_paragraph()
run = p.add_run('Before:')
run.bold = True
doc.add_paragraph('fund_requirements → agent EOA: threshold = 0.21 xDAI, topup = 2.0 xDAI', style='List Bullet')
doc.add_paragraph('auto_transfer: (does not exist)', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('After:')
run.bold = True
doc.add_paragraph('fund_requirements → agent EOA: threshold = 0.001 xDAI (safe_tx_value), topup = 0.01 xDAI', style='List Bullet')
doc.add_paragraph('auto_transfer → gnosis: threshold = 0.21 xDAI, topup = 2.0 xDAI', style='List Bullet')

doc.add_heading('Configuration Guideline', level=2)
p = doc.add_paragraph()
run = p.add_run('Deployment requirement: ')
run.bold = True
p.add_run(
    'The auto_transfer threshold must be configured higher than the gas cost of a Safe '
    'transaction on that chain. This ensures the EOA can always afford the auto-transfer '
    'when the threshold triggers. All current agent configurations meet this requirement.'
)

doc.add_page_break()

# ============================================================
# 8. IMPLEMENTATION PLAN
# ============================================================
doc.add_heading('8. Implementation Plan', level=1)

doc.add_heading('Phase 1: Core Logic (funds-manager repo)', level=2)
items = [
    'Add auto_transfer config to the Params model (backward compatible — defaults to empty).',
    'Create new FSM rounds: CheckFundsRound, PrepareSafeToEOATransferRound, FinishedFundsTransferRound.',
    'Create FSM behaviours tied to the rounds.',
    'Create FundsManagerAbciApp with transition mapping.',
    'Update skill.yaml to register new components and add TransactionSettlementAbci dependency.',
    'Existing get_funds_status() code and SimpleBehaviour — zero changes.',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Phase 2: Per-Agent Integration (external repos)', level=2)
items = [
    'Each agent updates composition.py to import and compose FundsManagerAbciApp at the start of their FSM cycle.',
    'Each agent updates service.yaml: lower EOA threshold to safe_tx_value, add auto_transfer config.',
    'Optimus: enable Safe native funding (currently disabled).',
    'Agents Fun: ensure FundsManagerAbciApp runs before CheckFundsBehaviour in the FSM.',
    'Polystrat: small handler fix to not suppress native POL deficit when needed for auto-transfer.',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Phase 3: Testing', level=2)
items = [
    'Unit tests for CheckFundsRound decision logic (all cases).',
    'Unit tests for Safe transaction payload building.',
    'Unit tests for FSM transitions (transfer vs no-transfer paths).',
    'Integration tests for full round-trip with mocked multicall and Safe tx.',
    'Per-agent tests for FSM composition.',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Phase 4: Follow-up', level=2)
doc.add_paragraph(
    'Address the Polystrat USDC→POL handler adjustment interaction. The handler must '
    'be made aware not to suppress native POL deficits that the auto-transfer depends on.'
)

doc.add_page_break()

# ============================================================
# 9. EDGE CASES & MITIGATIONS
# ============================================================
doc.add_heading('9. Edge Cases & Mitigations', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
add_table_row(table, ['Edge Case', 'What Happens', 'Mitigation'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Edge Case', 'What Happens', 'Mitigation'], header=True)

add_table_row(table, [
    'Auto-transfer depletes Safe',
    'Safe falls below its own threshold.',
    'Expected. Safe deficit reported to user via existing flow. User funds Safe.'
])
add_table_row(table, [
    'Transfer fails',
    'TransactionSettlementAbci retries. If all retries exhausted, no transfer this cycle.',
    'Next FSM cycle re-evaluates and retries. If EOA drains to safe_tx_value, Path 3 kicks in.'
])
add_table_row(table, [
    'Agent restart during transfer',
    'In-memory state lost.',
    'FSM reads fresh on-chain balances on restart. If transfer confirmed, balances reflect it. If not, re-attempts.'
])
add_table_row(table, [
    'Gas price spike',
    'EOA may not afford Safe tx at current gas price.',
    'CheckFundsRound estimates gas at runtime. Skips transfer if unaffordable. Retries next cycle.'
])
add_table_row(table, [
    'Concurrent Safe transactions',
    'Could conflict with other agent Safe txs.',
    'No conflict — FSM serializes all Safe transactions. Auto-transfer blocks other rounds while executing.'
])
add_table_row(table, [
    'Safe has only ERC20, no native',
    'Auto-transfer cannot move ERC20 tokens.',
    'Treated as "Safe empty." Safe deficit reported. User funds Safe with native tokens.'
])
add_table_row(table, [
    'Multiple chains',
    'Agent operates on multiple chains.',
    'Each chain evaluated independently. Transfers executed sequentially, one per chain.'
])
add_table_row(table, [
    'MW polls during in-flight transfer',
    'On-chain balances haven\'t updated yet.',
    'Benign. EOA is above safe_tx_value threshold, so no deficit reported. Transfer confirms shortly.'
])
style_table(table)

doc.add_page_break()

# ============================================================
# 10. DECISIONS MADE
# ============================================================
doc.add_heading('10. Decisions Made', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
add_table_row(table, ['Decision', 'Resolution', 'Rationale'], header=True)
table._tbl.remove(table.rows[0]._tr)
add_table_row(table, ['Decision', 'Resolution', 'Rationale'], header=True)

add_table_row(table, [
    'Single-agent or multi-agent?',
    'Single agent (1-of-1) only',
    'All current deployments are 1-of-1.'
])
add_table_row(table, [
    'Modify get_funds_status()?',
    'No — zero code changes',
    'Keep existing /funds-status interaction completely unchanged. Middleware and frontend need no updates.'
])
add_table_row(table, [
    'Transaction execution model?',
    'Dedicated FSM round',
    'The FSM serializes Safe transactions — no nonce conflicts. While auto-transfer executes, no other Safe tx can run.'
])
add_table_row(table, [
    'Where does the FSM live?',
    'In the funds-manager repo',
    'All agents import and compose the same AbciApp. No per-agent reimplementation.'
])
add_table_row(table, [
    'Separate safe_tx_gas_cost config?',
    'No',
    'Gas cost is a runtime value. Thresholds must be configured above it (deployment guideline).'
])
add_table_row(table, [
    'Optimus Safe enablement?',
    'Yes, required',
    'Optimus must enable Safe native funding to use auto-transfer.'
])
add_table_row(table, [
    'Move agent-specific swap logic?',
    'No — stays in each agent',
    'x402 swaps, wxDAI folding, USDC→POL conversion remain in per-agent handlers. Funds manager focuses on native token management.'
])
style_table(table)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/dhairya/Desktop/Work/Valory/Github/funds-manager/WD_Auto_Transfer_Safe_to_EOA.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
